"""Extraction du texte d'un CV déposé (PDF, DOCX, TXT, Markdown).

Le fichier reste la propriété de l'utilisateur : on n'en garde que le texte,
qui devient la matière première de la réécriture par offre.
"""

import io
import re
from dataclasses import dataclass, field

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".markdown")

# En dessous, on considère que l'extraction a échoué (PDF scanné, par exemple).
_MIN_USABLE_CHARS = 200


class UnsupportedFile(ValueError):
    """Format non pris en charge."""


@dataclass
class ExtractedCv:
    text: str
    pages: int = 0
    warnings: list[str] = field(default_factory=list)


def _clean(text: str) -> str:
    """Remet d'aplomb un texte sorti d'un PDF : césures, espaces, lignes vides."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\xa0", " ")
    # Césure de fin de ligne : « déve-\nloppeur » → « développeur »
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Espaces multiples, mais on garde les retours à la ligne
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r" +\n", "\n", text)
    # Trois lignes vides ou plus → une seule
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> ExtractedCv:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # fichier corrompu, chiffré…
        raise UnsupportedFile(f"PDF illisible ({type(exc).__name__})") from exc

    warnings: list[str] = []
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            raise UnsupportedFile("PDF protégé par mot de passe") from None
        warnings.append("PDF protégé : ouvert sans mot de passe, le texte peut être partiel.")

    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            warnings.append("Une page n'a pas pu être lue.")

    return ExtractedCv(text=_clean("\n\n".join(chunks)), pages=len(reader.pages), warnings=warnings)


def _from_docx(data: bytes) -> ExtractedCv:
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedFile(f"DOCX illisible ({type(exc).__name__})") from exc

    lines = [paragraph.text for paragraph in document.paragraphs]
    # Beaucoup de CV mettent leurs colonnes dans des tableaux invisibles.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" · ".join(cells))

    return ExtractedCv(text=_clean("\n".join(lines)))


def _from_text(data: bytes) -> ExtractedCv:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return ExtractedCv(text=_clean(data.decode(encoding)))
        except UnicodeDecodeError:
            continue
    raise UnsupportedFile("encodage de texte non reconnu")


def extract(filename: str, data: bytes) -> ExtractedCv:
    """Texte d'un CV déposé, quel que soit son format d'origine."""
    if not data:
        raise UnsupportedFile("fichier vide")

    name = (filename or "").lower().strip()
    if name.endswith(".pdf"):
        result = _from_pdf(data)
    elif name.endswith(".docx"):
        result = _from_docx(data)
    elif name.endswith((".txt", ".md", ".markdown")):
        result = _from_text(data)
    elif name.endswith(".doc"):
        raise UnsupportedFile(
            "l'ancien format .doc n'est pas pris en charge — enregistre en .docx ou .pdf"
        )
    else:
        raise UnsupportedFile(
            "format non pris en charge — formats acceptés : " + ", ".join(SUPPORTED_EXTENSIONS)
        )

    if len(result.text) < _MIN_USABLE_CHARS:
        result.warnings.append(
            "Très peu de texte extrait : si le CV est un scan ou une image, "
            "dépose une version texte (PDF exporté depuis Word, ou .docx)."
        )
    return result
