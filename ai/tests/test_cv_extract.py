"""Dépôt d'un CV : extraction du texte depuis PDF, DOCX, TXT, Markdown."""

import io

import pytest

from app.cv_extract import UnsupportedFile, extract

CV_TEXTE = """Camille Dupont
Développeuse Full Stack / DevOps — Nantes
camille.dupont@example.com

EXPÉRIENCES
Développeuse Full Stack — Studio Katana (2021 - aujourd'hui)
Application Node.js / React, intégration continue GitLab CI, déploiements Docker.

Développeuse back-end — Coopérative Bleuet (2019 - 2021)
API REST Node.js et MongoDB, revue de code et tests unitaires.

FORMATION
Master Informatique — Université de Nantes (2017 - 2019)
"""


def _docx_bytes(paragraphs: list[str]) -> bytes:
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _minimal_pdf(text: str) -> bytes:
    """PDF non compressé minimal, avec table xref valide (pypdf l'exige)."""
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>",
        b"<</Length %d>>stream\n%s\nendstream" % (len(content), content),
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj" % number + body + b"endobj\n"

    xref_at = len(out)
    out += b"xref\n0 %d\n" % (len(objects) + 1)
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += b"%010d 00000 n \n" % offset
    out += b"trailer<</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objects) + 1,
        xref_at,
    )
    return bytes(out)


def test_texte_brut():
    result = extract("cv.txt", CV_TEXTE.encode("utf-8"))

    assert "Camille Dupont" in result.text
    assert "Studio Katana" in result.text
    assert result.warnings == []


def test_markdown():
    result = extract("cv.md", b"# Camille Dupont\n\n" + CV_TEXTE.encode("utf-8"))

    assert result.text.startswith("# Camille Dupont")


def test_docx():
    data = _docx_bytes(CV_TEXTE.splitlines())
    result = extract("cv.docx", data)

    assert "Camille Dupont" in result.text
    assert "Master Informatique" in result.text


def test_docx_lit_aussi_les_tableaux():
    from docx import Document

    document = Document()
    document.add_paragraph("Camille Dupont")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "2021 - aujourd'hui"
    table.rows[0].cells[1].text = "Développeuse chez Studio Katana"
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract("cv.docx", buffer.getvalue())

    assert "Studio Katana" in result.text  # beaucoup de CV mettent tout en tableau


def test_pdf():
    result = extract("cv.pdf", _minimal_pdf("Camille Dupont - Developpeuse Full Stack"))

    assert "Camille Dupont" in result.text
    assert result.pages == 1


def test_pdf_sans_texte_previent():
    """Un CV scanné ne produit rien d'exploitable : il faut le dire."""
    result = extract("cv.pdf", _minimal_pdf(""))

    assert any("scan" in warning or "peu de texte" in warning for warning in result.warnings)


def test_cesures_et_espaces_nettoyes():
    result = extract("cv.txt", "déve-\nloppeuse    full     stack\n\n\n\nNantes".encode("utf-8"))

    assert "développeuse" in result.text
    assert "full stack" in result.text
    assert "\n\n\n" not in result.text


def test_ancien_format_doc_explique():
    with pytest.raises(UnsupportedFile, match="docx"):
        extract("cv.doc", b"contenu binaire")


def test_format_inconnu():
    with pytest.raises(UnsupportedFile, match="format non pris en charge"):
        extract("cv.pages", b"contenu")


def test_fichier_vide():
    with pytest.raises(UnsupportedFile, match="vide"):
        extract("cv.pdf", b"")


# --- Endpoint ----------------------------------------------------------


def test_endpoint_extract_cv(client):
    response = client.post(
        "/extract-cv",
        files={"file": ("cv.txt", CV_TEXTE.encode("utf-8"), "text/plain")},
    )
    assert response.status_code == 200

    body = response.json()
    assert "Camille Dupont" in body["text"]
    assert body["chars"] > 200
    assert body["filename"] == "cv.txt"


def test_endpoint_refuse_format_non_supporte(client):
    response = client.post(
        "/extract-cv",
        files={"file": ("cv.pages", b"contenu", "application/octet-stream")},
    )

    assert response.status_code == 415
    assert "format" in response.json()["detail"]


def test_endpoint_refuse_fichier_trop_lourd(client):
    response = client.post(
        "/extract-cv",
        files={"file": ("cv.pdf", b"x" * (6 * 1024 * 1024), "application/pdf")},
    )

    assert response.status_code == 413
